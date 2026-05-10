import os
import re
import markdown
import html2text
from docx import Document
from typing import Optional, List
from app.app_services.file_service import FileService
from app.app_services.pdf_conversion_service import PDFConversionService
from app.app_services.ebook_conversion_service import EBookConversionService
from app.app_core.exceptions import FileProcessingError

class MarkdownConversionService:
    """Service for handling Markdown conversions between various formats."""

    @staticmethod
    def pdf_to_markdown(input_path: str, output_path: str) -> str:
        """Convert PDF to Markdown format."""
        try:
            return PDFConversionService.pdf_to_markdown(input_path, output_path)
        except Exception as e:
            raise FileProcessingError(f"PDF to Markdown conversion failed: {str(e)}")

    @staticmethod
    def markdown_to_pdf(input_path: str, output_path: str) -> str:
        """Convert Markdown to PDF format."""
        try:
            return PDFConversionService.markdown_to_pdf(input_path, output_path)
        except Exception as e:
            raise FileProcessingError(f"Markdown to PDF conversion failed: {str(e)}")

    @staticmethod
    def markdown_to_html(input_path: str, output_path: str) -> str:
        """Convert Markdown to HTML format."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            html = markdown.markdown(content, extensions=['extra', 'codehilite', 'fenced_code', 'tables'])
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html)
            return output_path
        except Exception as e:
            raise FileProcessingError(f"Markdown to HTML conversion failed: {str(e)}")

    @staticmethod
    def markdown_to_epub(input_path: str, output_path: str) -> str:
        """Convert Markdown to ePUB format."""
        try:
            # EBookConversionService.markdown_to_epub returns the output path
            epub_path = EBookConversionService.markdown_to_epub(input_path)
            if epub_path != output_path:
                if os.path.exists(output_path):
                    os.remove(output_path)
                os.rename(epub_path, output_path)
            return output_path
        except Exception as e:
            raise FileProcessingError(f"Markdown to ePUB conversion failed: {str(e)}")

    @staticmethod
    def markdown_to_word(input_path: str, output_path: str) -> str:
        """Convert Markdown to Word document."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            doc = Document()
            for line in content.split('\n'):
                if line.strip().startswith('# '):
                   doc.add_heading(line.strip()[2:], level=1)
                elif line.strip().startswith('## '):
                   doc.add_heading(line.strip()[3:], level=2)
                else:
                   doc.add_paragraph(line)
            doc.save(output_path)
            return output_path
        except Exception as e:
            raise FileProcessingError(f"Markdown to Word conversion failed: {str(e)}")

    @staticmethod
    def markdown_to_latex(input_path: str, output_path: str) -> str:
        """Convert Markdown to LaTeX format (Simplified)."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Very basic LaTeX conversion
            latex = content
            latex = re.sub(r'^# (.*)$', r'\\section{\1}', latex, flags=re.MULTILINE)
            latex = re.sub(r'^## (.*)$', r'\\subsection{\1}', latex, flags=re.MULTILINE)
            latex = re.sub(r'\*\*(.*)\*\*', r'\\textbf{\1}', latex)
            latex = re.sub(r'\*(.*)\*', r'\\textit{\1}', latex)
            
            # Wrap in minimal document structure if it's not present
            if '\\documentclass' not in latex:
                latex = f"\\documentclass{{article}}\n\\begin{{document}}\n{latex}\n\\end{{document}}"
                
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(latex)
            return output_path
        except Exception as e:
            raise FileProcessingError(f"Markdown to LaTeX conversion failed: {str(e)}")

    @staticmethod
    def markdown_to_text(input_path: str, output_path: str) -> str:
        """Convert Markdown to plain text."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Strip most common markdown syntax
            text = content
            text = re.sub(r'#+\s+', '', text)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            text = re.sub(r'`+(.*?)`+', r'\1', text)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return output_path
        except Exception as e:
            raise FileProcessingError(f"Markdown to Text conversion failed: {str(e)}")

    @staticmethod
    def html_to_markdown(input_path: str, output_path: str) -> str:
        """Convert HTML to Markdown format."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            h = html2text.HTML2Text()
            h.ignore_links = False
            md = h.handle(content)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md)
            return output_path
        except Exception as e:
            raise FileProcessingError(f"HTML to Markdown conversion failed: {str(e)}")

    @staticmethod
    def word_to_markdown(input_path: str, output_path: str) -> str:
        """Convert Word document to Markdown format."""
        try:
            doc = Document(input_path)
            md_lines = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading 1'):
                    md_lines.append(f"# {para.text}")
                elif para.style.name.startswith('Heading 2'):
                    md_lines.append(f"## {para.text}")
                else:
                    md_lines.append(para.text)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(md_lines))
            return output_path
        except Exception as e:
            raise FileProcessingError(f"Word to Markdown conversion failed: {str(e)}")
